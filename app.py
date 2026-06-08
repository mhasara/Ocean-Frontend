import io
import os
import re
import docx
from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory

app = Flask(__name__)
app.secret_key = "sinhala_grammar_checker_secret_key"

# Configuration
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
DICTIONARY_PATH = os.path.join(BASE_DIR, "dictionary", "sinhala_words.txt")
ALLOWED_EXTENSIONS = {"txt", "docx"}

# Load dictionary
def load_dictionary():
    words = set()
    if os.path.exists(DICTIONARY_PATH):
        try:
            with open(DICTIONARY_PATH, "r", encoding="utf-8") as f:
                for line in f:
                    word = line.strip()
                    if word:
                        words.add(word)
        except Exception as e:
            print(f"Error loading dictionary: {e}")
    else:
        print(f"Dictionary file not found at: {DICTIONARY_PATH}")
    return words

# Initialize the dictionary set
dictionary_words = load_dictionary()

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def is_sinhala_word(token):
    # A word to check must contain at least one Sinhala character (\u0d80 to \u0dff)
    return any("\u0d80" <= char <= "\u0dff" for char in token)

@app.route("/", methods=["GET", "POST"])
def index():
    original_text = ""
    corrected_text_html = ""
    incorrect_words = []
    has_results = False
    error_message = None

    if request.method == "POST":
        # Check if file part is present
        if "file" not in request.files:
            error_message = "No file part in the request."
            return render_template(
                "index.html",
                original_text=original_text,
                corrected_text_html=corrected_text_html,
                incorrect_words=incorrect_words,
                has_results=has_results,
                error_message=error_message
            )

        file = request.files["file"]

        # If user does not select file, browser submits empty part without filename
        if file.filename == "":
            error_message = "No file selected. Please choose a .txt or .docx file to upload."
            return render_template(
                "index.html",
                original_text=original_text,
                corrected_text_html=corrected_text_html,
                incorrect_words=incorrect_words,
                has_results=has_results,
                error_message=error_message
            )

        if file and allowed_file(file.filename):
            try:
                filename = file.filename.lower()
                if filename.endswith(".docx"):
                    # Extract text from Word Document (.docx)
                    doc = docx.Document(io.BytesIO(file.read()))
                    text = "\n".join([p.text for p in doc.paragraphs])
                else:
                    # Read and decode text file (.txt)
                    file_bytes = file.read()
                    try:
                        text = file_bytes.decode("utf-8")
                    except UnicodeDecodeError:
                        # Fallback to utf-8-sig to handle UTF-8 with BOM, or latin-1 if all else fails
                        try:
                            text = file_bytes.decode("utf-8-sig")
                        except UnicodeDecodeError:
                            text = file_bytes.decode("latin-1")

                original_text = text
                
                # Split text into tokens: words and non-words (spaces, punctuation, etc.)
                # Sinhala Unicode range: U+0D80 to U+0DFF. ZWJ (U+200D) and ZWNJ (U+200C) are used in conjuncts.
                tokens = re.split(r"([^\u0d80-\u0dff\w\u200d\u200c]+)", text)
                
                corrected_tokens = []
                incorrect_set = set()

                # Refresh dictionary in case it was modified
                global dictionary_words
                dictionary_words = load_dictionary()

                for token in tokens:
                    if not token:
                        continue
                    
                    if is_sinhala_word(token):
                        # Clean word for comparison (strip any trailing/leading symbols just in case)
                        clean_word = token.strip()
                        
                        if clean_word in dictionary_words:
                            corrected_tokens.append(token)
                        else:
                            # Mark incorrect word
                            incorrect_set.add(clean_word)
                            # Wrap in HTML highlight tag
                            highlighted = f'<span class="incorrect-word" title="Not in dictionary">{token}</span>'
                            corrected_tokens.append(highlighted)
                    else:
                        # Non-Sinhala word or separator, keep as is
                        corrected_tokens.append(token)

                corrected_text_html = "".join(corrected_tokens)
                # Replace newline characters with HTML line breaks for rendering
                corrected_text_html = corrected_text_html.replace("\n", "<br>")
                
                incorrect_words = sorted(list(incorrect_set))
                has_results = True

            except Exception as e:
                error_message = f"Error processing file: {str(e)}"
        else:
            error_message = "Invalid file type. Only .txt and .docx files are allowed."

    return render_template(
        "index.html",
        original_text=original_text,
        corrected_text_html=corrected_text_html,
        incorrect_words=incorrect_words,
        has_results=has_results,
        error_message=error_message
    )

@app.route("/download-sample")
def download_sample():
    # Helper endpoint to download a sample Sinhala text file for testing
    static_dir = os.path.join(BASE_DIR, "static")
    sample_file_path = os.path.join(static_dir, "sample_sinhala.txt")
    
    # Create the sample file if it doesn't exist
    if not os.path.exists(sample_file_path):
        os.makedirs(static_dir, exist_ok=True)
        sample_content = (
            "මම ලංකාවේ ඉපදුණා. මම සිංහල ඉගෙන ගන්නවා.\n"
            "මගේ යාළුවා පාසලට ගියා.\n"
            "නමුත් ඔහු කනව (වැරදි - කනවා විය යුතුය).\n"
            "ළමයා සෙල්ලම් කරනව (වැරදි - කරනවා විය යුතුය).\n"
        )
        with open(sample_file_path, "w", encoding="utf-8") as f:
            f.write(sample_content)
            
    return send_from_directory(static_dir, "sample_sinhala.txt", as_attachment=True)

@app.route("/download-sample-docx")
def download_sample_docx():
    # Helper endpoint to download a sample Sinhala docx file for testing
    static_dir = os.path.join(BASE_DIR, "static")
    sample_file_path = os.path.join(static_dir, "sample_sinhala.docx")
    
    # Create the sample file if it doesn't exist (using python-docx)
    if not os.path.exists(sample_file_path):
        try:
            os.makedirs(static_dir, exist_ok=True)
            doc = docx.Document()
            doc.add_paragraph("මම ලංකාවේ ඉපදුණා. මම සිංහල ඉගෙන ගන්නවා.")
            doc.add_paragraph("මගේ යාළුවා පාසලට ගියා.")
            doc.add_paragraph("නමුත් ඔහු කනව (වැරදි - කනවා විය යුතුය).")
            doc.add_paragraph("ළමයා සෙල්ලම් කරනව (වැරදි - කරනවා විය යුතුය).")
            doc.save(sample_file_path)
        except Exception as e:
            print(f"Error creating docx: {e}")
            
    return send_from_directory(static_dir, "sample_sinhala.docx", as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True, port=5000)
